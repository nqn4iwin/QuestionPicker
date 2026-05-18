"""Export exam JSON to Word (.docx) with 2-column pages and intact question blocks."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

CHOICE_MARKERS = "①②③④"
DEFAULT_FONT = "맑은 고딕"
BODY_SIZE_PT = 10.5
COLUMN_WIDTH = Inches(3.15)
IMAGE_MAX_WIDTH = Inches(2.95)
# Per sheet: sort by question no., take 4, left col = lower pair, right col = upper pair (1|3, 2|4).
QUESTIONS_PER_WORD_PAGE = 4


def _choice_marker(choice_no: int) -> str:
    if 1 <= choice_no <= 4:
        return CHOICE_MARKERS[choice_no - 1]
    return f"{choice_no}."


def _set_run_font(run, *, name: str = DEFAULT_FONT, size_pt: float = BODY_SIZE_PT) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)


def _add_text_paragraph(
    cell,
    text: str,
    *,
    bold: bool = False,
    space_after_pt: float = 3,
) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    run = paragraph.add_run(text)
    run.bold = bold
    _set_run_font(run)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_table_borders_none(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def _set_cell_width(cell, width) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_pr.append(tc_w)


def _question_number(question: dict[str, Any]) -> int:
    try:
        return int(question.get("no", 0))
    except (TypeError, ValueError):
        return 0


def _sort_questions_by_number(
    questions: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    return sorted(questions, key=_question_number)


def _chunk_questions(
    questions: list[dict[str, Any]],
    *,
    chunk_size: int = QUESTIONS_PER_WORD_PAGE,
) -> list[list[dict[str, Any]]]:
    if chunk_size < 1:
        raise ValueError("chunk_size must be at least 1")
    return [
        questions[index : index + chunk_size]
        for index in range(0, len(questions), chunk_size)
    ]


def _is_search_result_document(document_data: dict[str, Any]) -> bool:
    search = document_data.get("search")
    if not isinstance(search, dict):
        return False
    matches = search.get("matches")
    return isinstance(matches, list) and bool(matches)


def _split_reading_order_columns(
    chunk: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Chunk must be sorted by question number (4 per sheet).

    Rows: lowest|third, second|highest — e.g. 1|3 then 2|4 for nos 1–4.
    """
    if not chunk:
        return [], []
    mid = (len(chunk) + 1) // 2
    return chunk[:mid], chunk[mid:]


def _resolve_asset_path(
    json_path: Path,
    meta: dict[str, Any],
    asset: dict[str, Any],
) -> Path | None:
    path_value = asset.get("path")
    if not isinstance(path_value, str) or not path_value.strip():
        return None
    raw = Path(path_value)
    if raw.is_file():
        return raw.resolve()
    base = json_path.parent
    candidate = base / raw
    if candidate.is_file():
        return candidate.resolve()
    assets_dir = meta.get("assets_dir")
    if isinstance(assets_dir, str) and assets_dir.strip():
        candidate = base / assets_dir / raw.name
        if candidate.is_file():
            return candidate.resolve()
    return None


def _add_question_block(
    cell,
    question: dict[str, Any],
    *,
    json_path: Path,
    meta: dict[str, Any],
) -> None:
    block = cell.add_table(rows=1, cols=1)
    _set_table_borders_none(block)
    block_cell = block.rows[0].cells[0]
    _set_cell_width(block_cell, COLUMN_WIDTH)

    number = question.get("no", "?")
    content = question.get("content", "")
    stem = f"{number}. {content}".strip() if content else f"{number}."
    _add_text_paragraph(block_cell, stem, bold=True, space_after_pt=4)

    assets = question.get("assets")
    if isinstance(assets, list):
        for asset in assets:
            if not isinstance(asset, dict) or asset.get("type") != "image":
                continue
            image_path = _resolve_asset_path(json_path, meta, asset)
            if image_path is None:
                continue
            paragraph = block_cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=IMAGE_MAX_WIDTH)

    choices = question.get("choices", [])
    if isinstance(choices, list):
        for choice in choices:
            if not isinstance(choice, dict):
                continue
            choice_no = int(choice.get("choice", 0))
            text = choice.get("text", "")
            if not isinstance(text, str) or not text.strip():
                continue
            line = f"{_choice_marker(choice_no)} {text.strip()}"
            _add_text_paragraph(block_cell, line, space_after_pt=2)

    spacer = block_cell.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    _set_row_cant_split(block.rows[0])


def _fill_cell_question_stack(
    cell,
    questions: list[dict[str, Any]],
    *,
    json_path: Path,
    meta: dict[str, Any],
) -> None:
    """Stack questions top-to-bottom in one column cell (read down, then across)."""
    for index, question in enumerate(questions):
        if index > 0:
            gap = cell.add_paragraph()
            gap.paragraph_format.space_after = Pt(6)
        _add_question_block(cell, question, json_path=json_path, meta=meta)


def _add_exam_sheet_table(
    document: Document,
    rows: list[tuple[dict[str, Any] | None, dict[str, Any] | None]],
    *,
    json_path: Path,
    meta: dict[str, Any],
) -> None:
    """Full exam: one question per side per row (1|3, then 2|4)."""
    if not rows:
        return

    page_table = document.add_table(rows=len(rows), cols=2)
    _set_table_borders_none(page_table)

    for row_index, (left_q, right_q) in enumerate(rows):
        row = page_table.rows[row_index]
        _set_row_cant_split(row)
        left_cell, right_cell = row.cells[0], row.cells[1]
        _set_cell_width(left_cell, COLUMN_WIDTH)
        _set_cell_width(right_cell, COLUMN_WIDTH)
        if left_q is not None:
            _add_question_block(
                left_cell, left_q, json_path=json_path, meta=meta
            )
        if right_q is not None:
            _add_question_block(
                right_cell, right_q, json_path=json_path, meta=meta
            )

    trailing = document.add_paragraph()
    trailing.paragraph_format.space_after = Pt(0)


def _add_search_result_table(
    document: Document,
    questions: list[dict[str, Any]],
    *,
    json_path: Path,
    meta: dict[str, Any],
    group_size: int = QUESTIONS_PER_WORD_PAGE,
) -> None:
    """Search hits: one table row per 4-number group; stack in each column cell."""
    chunks = _chunk_questions(questions, chunk_size=group_size)
    if not chunks:
        return

    page_table = document.add_table(rows=len(chunks), cols=2)
    _set_table_borders_none(page_table)

    for row_index, chunk in enumerate(chunks):
        left, right = _split_reading_order_columns(chunk)
        row = page_table.rows[row_index]
        left_cell, right_cell = row.cells[0], row.cells[1]
        _set_cell_width(left_cell, COLUMN_WIDTH)
        _set_cell_width(right_cell, COLUMN_WIDTH)
        if left:
            _fill_cell_question_stack(
                left_cell, left, json_path=json_path, meta=meta
            )
        if right:
            _fill_cell_question_stack(
                right_cell, right, json_path=json_path, meta=meta
            )

    trailing = document.add_paragraph()
    trailing.paragraph_format.space_after = Pt(0)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def json_to_docx(
    json_path: Path | str,
    output_path: Path | str | None = None,
    *,
    questions_per_page: int = QUESTIONS_PER_WORD_PAGE,
) -> Path:
    """Write exam questions to a 2-column Word document."""
    source = Path(json_path).resolve()
    if not source.is_file():
        raise FileNotFoundError(source)

    document_data = json.loads(source.read_text(encoding="utf-8"))
    questions = document_data.get("questions", [])
    if not isinstance(questions, list) or not questions:
        raise ValueError("JSON has no questions to export")

    meta = document_data.get("meta", {})
    if not isinstance(meta, dict):
        meta = {}

    valid_questions = [q for q in questions if isinstance(q, dict)]
    ordered = _sort_questions_by_number(valid_questions)
    is_search = _is_search_result_document(document_data)

    document = Document()
    _configure_document(document)

    if is_search:
        _add_search_result_table(
            document,
            ordered,
            json_path=source,
            meta=meta,
            group_size=questions_per_page,
        )
    else:
        page_chunks = _chunk_questions(ordered, chunk_size=questions_per_page)
        for index, chunk in enumerate(page_chunks):
            if index > 0:
                break_paragraph = document.add_paragraph()
                break_paragraph.add_run().add_break(WD_BREAK.PAGE)

            left, right = _split_reading_order_columns(chunk)
            chunk_rows = [
                (
                    left[i] if i < len(left) else None,
                    right[i] if i < len(right) else None,
                )
                for i in range(max(len(left), len(right)))
            ]
            _add_exam_sheet_table(
                document, chunk_rows, json_path=source, meta=meta
            )

    dest = Path(output_path) if output_path else source.with_suffix(".docx")
    dest = dest.resolve()
    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(dest)
    return dest


def json_to_docx_and_pdf(
    json_path: Path | str,
    docx_path: Path | str | None = None,
    pdf_path: Path | str | None = None,
    *,
    questions_per_page: int = QUESTIONS_PER_WORD_PAGE,
) -> tuple[Path, Path]:
    """Write .docx then convert to .pdf (Windows + MS Word required for PDF)."""
    from modules.docx_to_pdf import convert_docx_to_pdf
    from modules.output_paths import default_pdf_path

    docx = json_to_docx(
        json_path,
        docx_path,
        questions_per_page=questions_per_page,
    )
    pdf_dest = (
        Path(pdf_path).resolve()
        if pdf_path is not None
        else default_pdf_path(docx_path=docx)
    )
    pdf = convert_docx_to_pdf(docx, pdf_dest)
    return docx, pdf
